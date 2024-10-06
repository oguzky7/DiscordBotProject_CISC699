import os
from fpdf import FPDF
from docx import Document

# Directory where the project files are located
directory = r"D:\HARRISBURG\Harrisburg Master's Fifth Term Late Summer\CISC 699\DiscordBotProject_CISC699"

# Lists for files and folders to ignore
files_to_ignore = ['ignore_this.py', '*txt', '*md', '*.pdf', '*.docx', '*.pyc', 'Config.py']  # Example file names to ignore
folders_to_ignore = ['ignore_folder', '.git', '__pycache__', 'PersonelTest', 'MockTesting', 'ExportedFiles', 'other', 'UnitTesting']  # Folders to ignore

def create_pdf(text_blocks, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    for title, content in text_blocks:
        pdf.add_page()  # Ensure each file starts on a new page
        pdf.multi_cell(0, 10, title)
        for line in content.split("\n"):
            pdf.multi_cell(0, 10, line.encode('latin1', 'replace').decode('latin1'))

    pdf.output(output_path)

def create_docx(text_blocks, output_path):
    doc = Document()
    for title, content in text_blocks:
        if doc.paragraphs:
            doc.add_page_break()  # Ensure each file starts on a new page
        doc.add_paragraph(title)
        doc.add_paragraph(content)
    doc.save(output_path)

def extract_project_text(directory, ignore_files=None, ignore_folders=None):
    if ignore_files is None:
        ignore_files = []
    if ignore_folders is None:
        ignore_folders = []

    text_blocks = []
    for root, dirs, files in os.walk(directory):
        dirs[:] = [d for d in dirs if d not in ignore_folders]
        
        for file in files:
            if file in ignore_files or not file.endswith('.py'):
                continue

            file_path = os.path.join(root, file)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    file_content = f.read()
                    text_blocks.append((f"--- {file} ---", file_content))
            except Exception as e:
                print(f"Could not read file {file_path}: {e}")

    return text_blocks

def create_documentation(directory, ignore_files=None, ignore_folders=None):
    text_blocks = extract_project_text(directory, ignore_files, ignore_folders)
    if text_blocks:
        pdf_path = os.path.join(directory, "project_text.pdf")
        docx_path = os.path.join(directory, "project_text.docx")
        create_pdf(text_blocks, pdf_path)
        create_docx(text_blocks, docx_path)
        print(f"Documentation created at: {pdf_path} and {docx_path}")
    else:
        print("No project text found.")

# Example of running the documentation creation
create_documentation(directory, files_to_ignore, folders_to_ignore)
